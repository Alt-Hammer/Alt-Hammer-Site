"""
Heading-driven extraction of Stratagems from the Core Rules Word document.

This script expects the Word source to use headings (H6 for stratagem name,
H7 subsections for metadata). It preserves inline formatting (bold/italic)
and a character style containing 'Keyword' by emitting minimal HTML for runs.

Output: src/data/stratagems.ts
"""

import sys
import os
import re
import html
from docx import Document

# Configuration
# Source document paths are centralised in source_paths.py — see that file
# to change where the documents live or what they are called.
from source_paths import CORE_RULES_DOCX
OUTPUT_PATH = os.path.join("src", "data", "stratagems.ts")


def sanitize_text(text: str) -> str:
    if text is None:
        return ""
    replacements = {
        '\u2013': '-', '\u2014': '-', '\u2018': "'", '\u2019': "'",
        '\u201c': '"', '\u201d': '"', '\u2026': '...', '\u2022': '*',
        '\u00b0': ' degrees', '\u2122': '(TM)', '\u00ae': '(R)', '\u00a9': '(C)'
    }
    for old, new in replacements.items():
        text = text.replace(old, new)
    return text


def slugify(text: str) -> str:
    text = (text or '').lower().strip()
    text = re.sub(r"[^\w\s-]", '', text)
    text = re.sub(r"[\s_]+", '-', text)
    text = re.sub(r"-+", '-', text)
    return text.strip('-')


def paragraph_heading_level(para):
    if not para.style or not getattr(para.style, 'name', None):
        return None
    m = re.search(r'heading\s*(\d+)', para.style.name.lower())
    return int(m.group(1)) if m else None


def find_section_by_heading(doc, heading_text: str):
    for para in doc.paragraphs:
        lvl = paragraph_heading_level(para)
        if lvl is not None and heading_text.lower() in para.text.lower():
            return para
    return None


def get_section_paragraphs(doc, parent_heading: str, subsection_text: str):
    parent = find_section_by_heading(doc, parent_heading)
    if not parent:
        return []
    parent_level = paragraph_heading_level(parent) or 0
    parent_index = None
    for i, p in enumerate(doc.paragraphs):
        if getattr(p, '_p', None) is getattr(parent, '_p', None):
            parent_index = i
            break
    if parent_index is None:
        return []

    subsection = None
    for para in doc.paragraphs[parent_index + 1:]:
        lvl = paragraph_heading_level(para)
        if lvl is not None and lvl <= parent_level:
            break
        if lvl is not None and subsection_text.lower() in para.text.lower():
            subsection = para
            break
    if not subsection:
        return []

    subsection_level = paragraph_heading_level(subsection) or 0
    paras = []
    for para in doc.paragraphs[get_paragraph_index(doc, subsection) + 1:]:
        lvl = paragraph_heading_level(para)
        if lvl is not None and lvl <= subsection_level:
            break
        paras.append(para)
    return paras


def get_paragraph_index(doc, target_para):
    for i, p in enumerate(doc.paragraphs):
        if getattr(p, '_p', None) is getattr(target_para, '_p', None):
            return i
    return None


def para_to_html(para):
    parts = []
    for run in para.runs:
        txt = sanitize_text(run.text or '')
        if not txt:
            continue
        txt = html.escape(txt)
        content = txt
        style_name = ''
        try:
            style_name = (run.style.name or '').lower() if run.style else ''
        except Exception:
            style_name = ''
        if getattr(run, 'bold', False):
            content = f"<strong>{content}</strong>"
        if getattr(run, 'italic', False):
            content = f"<em>{content}</em>"
        if style_name and 'keyword' in style_name:
            content = f"<span class=\"keyword\">{content}</span>"
        parts.append(content)
    return ''.join(parts)


def extract_stratagems_from_headings(paragraphs) -> list:
    stratagems = []
    current = None
    current_sub = None

    for para in paragraphs:
        lvl = paragraph_heading_level(para)
        plain = sanitize_text(para.text.strip())
        if not plain:
            continue

        if lvl == 6:
            if current:
                stratagems.append(current)
            current = {
                'id': slugify(plain),
                'name': plain,
                'maxUsesPerRound': '',
                'timing': '',
                'description': '',
                'cost': '',
            }
            current_sub = None
            continue

        if lvl == 7 and current:
            heading_label = plain.lower()
            current_sub = None
            if (('when' in heading_label and 'use' in heading_label) or 'when stratagem' in heading_label or 'when used' in heading_label):
                current_sub = 'timing'
            elif 'description' in heading_label:
                current_sub = 'description'
            elif 'cost' in heading_label or 'command point' in heading_label:
                current_sub = 'cost'
            elif 'max uses' in heading_label or 'max uses per battle round' in heading_label:
                current_sub = 'maxUsesPerRound'
            else:
                current_sub = None
            continue

        if current and current_sub:
            html_val = para_to_html(para)
            if current.get(current_sub):
                current[current_sub] += '\n' + html_val
            else:
                current[current_sub] = html_val

    if current:
        stratagems.append(current)
    return stratagems


def generate_typescript_file(stratagems: list) -> str:
    interface = """export interface Stratagem {
  id: string;
  name: string;
  maxUsesPerRound?: string;
  timing: string;
  description: string;
  cost: string;
  subRows?: Array<{
    effect: string;
    cost: string;
  }>;
}
"""
    data_lines = ["export const STRATAGEMS: Stratagem[] = ["]
    for s in stratagems:
        data_lines.append("  {")
        data_lines.append(f'    id: "{s["id"]!s}",')
        data_lines.append(f'    name: "{escape_string(s["name"])}",')
        data_lines.append(f'    maxUsesPerRound: "{escape_string(s.get("maxUsesPerRound", ""))}",')
        data_lines.append(f'    timing: "{escape_string(s.get("timing", ""))}",')
        data_lines.append(f'    description: "{escape_string(s.get("description", ""))}",')
        data_lines.append(f'    cost: "{escape_string(s.get("cost", ""))}",')
        if s.get('subRows'):
            data_lines.append('    subRows: [')
            for sub in s['subRows']:
                data_lines.append('      {')
                data_lines.append(f'        effect: "{escape_string(sub.get("effect", ""))}",')
                data_lines.append(f'        cost: "{escape_string(sub.get("cost", ""))}",')
                data_lines.append('      },')
            data_lines.append('    ],')
        data_lines.append('  },')
    data_lines.append('];')
    header = """// Auto-generated by extract_stratagems.py
// DO NOT EDIT MANUALLY — regenerate from Word document

"""
    return header + interface + "\n" + "\n".join(data_lines)


def escape_string(s: str) -> str:
    s = str(s or '')
    s = s.replace('\\', '\\\\')
    s = s.replace('"', '\\"')
    s = s.replace('\n', '\\n')
    s = s.replace('\r', '\\r')
    s = s.replace('\t', '\\t')
    return s


def extract_stratagems(docx_path: str = CORE_RULES_DOCX, output_path: str = OUTPUT_PATH) -> bool:
    print('\n' + '=' * 60)
    print('  Extracting Stratagems from Core Rules')
    print('=' * 60)
    print('  Source: ', docx_path)
    print('  Output: ', output_path)
    print('=' * 60 + '\n')

    if not os.path.exists(docx_path):
        print('  ✗  ERROR: Source file not found: ', docx_path)
        return False

    doc = Document(docx_path)
    paras = get_section_paragraphs(doc, 'Command Points & Stratagems', 'Stratagems')
    if not paras:
        print("  ✗  ERROR: Could not find 'Stratagems' subsection")
        return False

    stratagems = extract_stratagems_from_headings(paras)
    if not stratagems:
        print('  ✗  ERROR: No stratagems found in headings')
        return False

    ts = generate_typescript_file(stratagems)
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write(ts)
    print(f'  ✓  Wrote {len(stratagems)} stratagems to {output_path}')
    return True


if __name__ == '__main__':
    docx_path = sys.argv[1] if len(sys.argv) > 1 else CORE_RULES_DOCX
    out_path = sys.argv[2] if len(sys.argv) > 2 else OUTPUT_PATH
    ok = extract_stratagems(docx_path, out_path)
    sys.exit(0 if ok else 1)
