from pathlib import Path
from docx import Document
import sys

sys.path.insert(0, str(Path(__file__).resolve().parent))
from extract_objectives import paragraph_heading_level, find_section_by_heading
from extract_stratagems import find_section_by_heading as find_section_by_heading_strat

doc_path = Path(r'C:/Users/alexc/OneDrive/04 Documents/Warhammer 40k/Alt-Hammer Standalone/Alt-Hammer 40,000 1st Edition - Core Rules.docx')
doc = Document(doc_path)

parent = find_section_by_heading(doc, 'Generating a Battle')
print('parent OBJ', getattr(parent, 'text', None), getattr(parent.style, 'name', None), paragraph_heading_level(parent))
print('parent in paragraphs?', parent in doc.paragraphs)
try:
    print('parent index', doc.paragraphs.index(parent))
except ValueError:
    print('parent not found by index()')
print('parent id', id(parent), 'len paragraphs', len(doc.paragraphs))
print('doc.paragraphs type', type(doc.paragraphs))
print('paragraph[787]', type(doc.paragraphs[787]), repr(doc.paragraphs[787].text), getattr(doc.paragraphs[787].style, 'name', None), paragraph_heading_level(doc.paragraphs[787]))
print('parent text equal?', parent.text == doc.paragraphs[787].text)
print('parent style equal?', getattr(parent.style, 'name', None), getattr(doc.paragraphs[787].style, 'name', None))
print('parent le text?', parent.text.lower() == doc.paragraphs[787].text.lower())
parent_level = paragraph_heading_level(parent) or 0

found_parent = False
for i, para in enumerate(doc.paragraphs):
    if para == parent:
        found_parent = True
        print('found parent at', i, 'id', id(para), 'text', repr(para.text))
        continue
    if not found_parent:
        continue
    level = paragraph_heading_level(para)
    text = para.text.strip()
    print('i', i, 'level', level, 'text', repr(text))
    if level is not None and level <= parent_level:
        print('break at', i, level, text)
        break
    if level is not None and 'optional game feature: secondary mission objectives'.lower() in text.lower():
        print('match at', i, level, text)
        break

print('\nStratagems parent:')
parent2 = find_section_by_heading_strat(doc, 'Command Points & Stratagems')
print('parent STRAT', getattr(parent2, 'text', None), getattr(parent2.style, 'name', None), paragraph_heading_level(parent2))

found_parent = False
for i, para in enumerate(doc.paragraphs):
    if para == parent2:
        found_parent = True
        print('found strat parent at', i)
        continue
    if not found_parent:
        continue
    level = paragraph_heading_level(para)
    text = para.text.strip()
    print('i', i, 'level', level, 'text', repr(text))
    if level is not None and level <= paragraph_heading_level(parent2):
        print('break2 at', i, level, text)
        break
    if level is not None and 'stratagems'.lower() in text.lower():
        print('match2 at', i, level, text)
        break
