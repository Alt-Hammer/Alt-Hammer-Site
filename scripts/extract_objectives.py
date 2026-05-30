"""
extract_objectives.py
─────────────────────
Reads Secondary Mission Objectives from the Alt-Hammer Core Rules Word document
and writes them as a structured TypeScript file consumed by ObjectiveGrid.astro.

WHAT IT DOES
────────────
1. Parses the Core Rules .docx file
2. Finds the "Generating a Battle" section
3. Extracts the "Optional Game Feature: Secondary Mission Objectives" subsection
4. Converts table data to structured SecondaryObjective objects with typed fields
5. Writes src/data/objectives.ts

OUTPUT
──────
  src/data/objectives.ts

File structure:
  export interface ObjectivePoints {
    condition: string;
    points: number;
  }

  export interface SecondaryObjective {
    id: string;
    number: number;
    name: string;
    timing: string;
    description: string;
    pointsAwarded: ObjectivePoints[];
    gameEffects?: string;
  }

  export const SECONDARY_OBJECTIVES: SecondaryObjective[] = [ ... ];

HOW TO RUN
──────────
From the alt-hammer-site project folder:
  python scripts/extract_objectives.py

Or use run_all.py to run it as part of the full pipeline.

CONFIGURATION
─────────────
Edit the paths below if your file locations change.
"""

import sys
import os
import re
from docx import Document

# ── Path configuration ────────────────────────────────────────────────────────

# Absolute path to your Core Rules Word document
CORE_RULES_DOCX = r"C:\Users\alexc\OneDrive\04 Documents\Warhammer 40k\Alt-Hammer Standalone\Alt-Hammer 40,000 1st Edition - Core Rules.docx"

# Output path
OUTPUT_PATH = os.path.join("src", "data", "objectives.ts")

# ─────────────────────────────────────────────────────────────────────────────

def sanitize_text(text: str) -> str:
    """Replace special Unicode characters with ASCII equivalents."""
    replacements = {
        '\u2013': '-',        # en dash
        '\u2014': '-',        # em dash
        '\u2018': "'",        # left single quote
        '\u2019': "'",        # right single quote
        '\u201c': '"',        # left double quote
        '\u201d': '"',        # right double quote
        '\u2026': '...',      # ellipsis
        '\u2022': '*',        # bullet
        '\u00b0': ' degrees', # degree symbol
        '\u2122': '(TM)',     # trademark
        '\u00ae': '(R)',      # registered
        '\u00a9': '(C)',      # copyright
    }
    for old, new in replacements.items():
        text = text.replace(old, new)
    return text


def slugify(text: str) -> str:
    """Convert text to kebab-case slug."""
    text = text.lower().strip()
    text = re.sub(r'[^\w\s-]', '', text)
    text = re.sub(r'[\s_]+', '-', text)
    text = re.sub(r'-+', '-', text)
    return text.strip('-')


def extract_cell_text(cell) -> str:
    """Extract all text from a table cell."""
    text_parts = []
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            text_parts.append(sanitize_text(run.text))
    return ''.join(text_parts).strip()


def paragraph_heading_level(para):
    """Return the numeric heading level for a paragraph, or None if it is not a heading."""
    if not para.style or not para.style.name:
        return None

    match = re.search(r'heading\s*(\d+)', para.style.name.lower())
    if match:
        return int(match.group(1))
    return None


def get_paragraph_index(doc, target_para):
    """Return the index of a paragraph in the document using the underlying XML element."""
    for index, para in enumerate(doc.paragraphs):
        if getattr(para, '_p', None) is getattr(target_para, '_p', None):
            return index
    return None


def log_section_headings(doc, parent_heading: str, limit: int = 20):
    """Print heading candidates under a parent section for debugging."""
    parent = find_section_by_heading(doc, parent_heading)
    if not parent:
        return

    parent_index = get_paragraph_index(doc, parent)
    if parent_index is None:
        return

    count = 0
    for i, para in enumerate(doc.paragraphs[parent_index + 1:], start=parent_index + 1):
        level = paragraph_heading_level(para)
        if level is not None:
            print(f"    [{i}] {para.style.name} | {para.text}")
            count += 1
            if count >= limit:
                break


def find_section_by_heading(doc, heading_text: str):
    """Find the first heading whose text matches heading_text."""
    for para in doc.paragraphs:
        level = paragraph_heading_level(para)
        if level is not None and heading_text.lower() in para.text.lower():
            return para
    return None


def find_subsection_by_heading(doc, parent_heading: str, subsection_text: str):
    """Find a subsection heading after a parent section heading."""
    parent = find_section_by_heading(doc, parent_heading)
    if not parent:
        return None

    parent_level = paragraph_heading_level(parent) or 0
    parent_index = get_paragraph_index(doc, parent)
    if parent_index is None:
        return None

    for para in doc.paragraphs[parent_index + 1:]:
        level = paragraph_heading_level(para)
        if level is not None and level <= parent_level:
            break

        if level is not None and subsection_text.lower() in para.text.lower():
            return para

    return None


def get_section_paragraphs(doc, parent_heading: str, subsection_text: str):
    """Return all paragraphs in a subsection until the next sibling heading."""
    subsection = find_subsection_by_heading(doc, parent_heading, subsection_text)
    if not subsection:
        return []

    subsection_level = paragraph_heading_level(subsection) or 0
    subsection_index = get_paragraph_index(doc, subsection)
    if subsection_index is None:
        return []

    paragraphs = []
    for para in doc.paragraphs[subsection_index + 1:]:
        level = paragraph_heading_level(para)
        if level is not None and level <= subsection_level:
            break

        paragraphs.append(para)

    return paragraphs


def find_table_after_heading(doc, heading_para):
    """Find the first table that appears after a given paragraph."""
    if not heading_para:
        return None
    
    heading_index = get_paragraph_index(doc, heading_para)
    if heading_index is None:
        return None

    found = False
    
    for i, element in enumerate(doc.element.body):
        if found and element.tag.endswith('tbl'):
            # Convert element to Table object
            from docx.table import Table
            return Table(element, doc)
        
        if element.tag.endswith('p'):
            para_index = sum(1 for e in doc.element.body[:i+1] if e.tag.endswith('p')) - 1
            if para_index == heading_index:
                found = True
    
    return None


def parse_points_cell(points_text: str) -> list:
    """
    Parse a points awarded cell which may contain multiple conditions.
    Examples:
      "1 per unit 3 per Battleline unit"
      "5"
      "2 per unit 4 per Battleline unit"
    
    Returns list of dicts: [{'condition': 'per unit', 'points': 1}, ...]
    """
    points = []
    
    # Clean up text
    points_text = points_text.strip()
    
    # Try to find all "number per X" patterns
    pattern = r'(\d+)\s+(?:per\s+)?([^0-9\n]+?)(?=\d+\s+|$)'
    matches = re.findall(pattern, points_text)
    
    if matches:
        for match in matches:
            num_str = match[0].strip()
            condition = match[1].strip()
            try:
                num = int(num_str)
                points.append({
                    'condition': condition,
                    'points': num,
                })
            except ValueError:
                pass
    else:
        # If pattern doesn't match, try simple number extraction
        try:
            # Look for just a number
            num_match = re.search(r'^\d+', points_text)
            if num_match:
                num = int(num_match.group())
                points.append({
                    'condition': 'base',
                    'points': num,
                })
        except ValueError:
            pass
    
    return points


def extract_objectives_from_table(table) -> list:
    """
    Parse secondary objectives table.
    
    Table structure (6 columns):
      Row 0: Header (# | Name | Completion Timing | Description | Points Awarded | Game Effects)
      Rows 1+: Data rows
      Note: Many rows have multiple condition rows (same name, different conditions)
    """
    objectives = []
    current_objective = None
    
    for row_idx, row in enumerate(table.rows):
        # Skip header row
        if row_idx == 0:
            continue
        
        cells = [extract_cell_text(cell) for cell in row.cells]
        
        # Skip completely empty rows
        if not any(cells):
            continue
        
        number_str = cells[0].strip() if len(cells) > 0 else ""
        name = cells[1].strip() if len(cells) > 1 else ""
        timing = cells[2].strip() if len(cells) > 2 else ""
        description = cells[3].strip() if len(cells) > 3 else ""
        points_text = cells[4].strip() if len(cells) > 4 else ""
        effects = cells[5].strip() if len(cells) > 5 else ""
        
        # Try to parse the number
        try:
            number = int(number_str) if number_str else 0
        except ValueError:
            number = 0
        
        # Determine if this is a new objective or a continuation
        if number_str and name:
            # This is a new objective
            points = parse_points_cell(points_text)
            
            current_objective = {
                'id': slugify(name),
                'number': number,
                'name': name,
                'timing': timing,
                'description': description,
                'pointsAwarded': points,
            }
            
            if effects:
                current_objective['gameEffects'] = effects
            
            objectives.append(current_objective)
        
        elif points_text and current_objective:
            # This is a continuation row with additional point conditions
            additional_points = parse_points_cell(points_text)
            current_objective['pointsAwarded'].extend(additional_points)
            
            if effects and not current_objective.get('gameEffects'):
                current_objective['gameEffects'] = effects
    
    return objectives


def extract_objectives_from_headings(paragraphs) -> list:
    """Parse objectives from a heading-based subsection."""
    objectives = []
    current_objective = None
    current_subheading = None

    for para in paragraphs:
        level = paragraph_heading_level(para)
        text = sanitize_text(para.text.strip())
        if not text:
            continue

        if level == 6:
            if current_objective:
                objectives.append(current_objective)

            current_objective = {
                'id': slugify(text),
                'number': 0,
                'name': text,
                'timing': '',
                'description': '',
                'pointsAwarded': [],
            }
            current_subheading = None
            continue

        if level == 7 and current_objective:
            heading_label = text.lower()
            current_subheading = None
            if 'when used' in heading_label:
                current_subheading = 'timing'
                current_objective['timing'] = ''
            elif 'description' in heading_label:
                current_subheading = 'description'
                current_objective['description'] = ''
            elif 'cost' in heading_label or 'points awarded' in heading_label:
                current_subheading = 'pointsAwarded'
            else:
                current_subheading = None
            continue

        if current_objective and current_subheading:
            if current_subheading == 'pointsAwarded':
                lines = [line.strip() for line in text.splitlines() if line.strip()]
                for line in lines:
                    points = parse_points_cell(line)
                    if points:
                        current_objective['pointsAwarded'].extend(points)
            else:
                if current_objective[current_subheading]:
                    current_objective[current_subheading] += '\n' + text
                else:
                    current_objective[current_subheading] = text

    if current_objective:
        objectives.append(current_objective)

    return objectives


def extract_objectives(docx_path: str = CORE_RULES_DOCX, output_path: str = OUTPUT_PATH) -> None:
    """
    Main extraction function.
    """
    print(f"\n{'='*60}")
    print(f"  Extracting Secondary Mission Objectives from Core Rules")
    print(f"{'='*60}")
    print(f"  Source:  {docx_path}")
    print(f"  Output:  {output_path}")
    print(f"{'='*60}\n")
    
    # Check source file exists
    if not os.path.exists(docx_path):
        print(f"  ✗  ERROR: Source file not found: {docx_path}")
        return False
    
    try:
        doc = Document(docx_path)
        
        # Find the objectives subsection
        objectives_paragraphs = get_section_paragraphs(
            doc,
            "Generating a Battle",
            "Optional Game Feature: Secondary Mission Objectives"
        )
        
        if not objectives_paragraphs:
            print("  ✗  ERROR: Could not find 'Secondary Mission Objectives' subsection")
            print("  DEBUG: Available headings under 'Generating a Battle':")
            log_section_headings(doc, "Generating a Battle")
            return False
        
        # Extract objective data from heading structure first
        objectives = extract_objectives_from_headings(objectives_paragraphs)
        
        if not objectives:
            # Fallback to table extraction if headings are not present
            objectives_heading = find_subsection_by_heading(
                doc,
                "Generating a Battle",
                "Optional Game Feature: Secondary Mission Objectives"
            )
            table = find_table_after_heading(doc, objectives_heading)
            
            if not table:
                print("  ✗  ERROR: Could not find Objectives table")
                return False
            
            objectives = extract_objectives_from_table(table)
            
            if not objectives:
                print("  ✗  ERROR: No objectives found in table")
                return False
        
        print(f"  Found {len(objectives)} objectives\n")
        
        # Generate TypeScript file
        ts_content = generate_typescript_file(objectives)
        
        # Write output
        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(ts_content)
        
        print(f"  ✓  Wrote {len(objectives)} objectives to {output_path}")
        print(f"{'='*60}\n")
        return True
    
    except Exception as e:
        print(f"  ✗  ERROR: {e}")
        import traceback
        traceback.print_exc()
        return False


def generate_typescript_file(objectives: list) -> str:
    """Generate the TypeScript source code for objectives."""
    
    # Build the interfaces
    interfaces = """export interface ObjectivePoints {
  condition: string;
  points: number;
}

export interface SecondaryObjective {
  id: string;
  number: number;
  name: string;
  timing: string;
  description: string;
  pointsAwarded: ObjectivePoints[];
  gameEffects?: string;
}
"""
    
    # Build the data array
    data_lines = ["export const SECONDARY_OBJECTIVES: SecondaryObjective[] = ["]
    
    for obj in objectives:
        data_lines.append("  {")
        data_lines.append(f'    id: "{obj["id"]!s}",')
        data_lines.append(f'    number: {obj["number"]!s},')
        data_lines.append(f'    name: "{escape_string(obj["name"])}",')
        data_lines.append(f'    timing: "{escape_string(obj["timing"])}",')
        data_lines.append(f'    description: "{escape_string(obj["description"])}",')
        
        data_lines.append("    pointsAwarded: [")
        for point in obj['pointsAwarded']:
            data_lines.append("      {")
            data_lines.append(f'        condition: "{escape_string(point["condition"])}",')
            data_lines.append(f'        points: {point["points"]!s},')
            data_lines.append("      },")
        data_lines.append("    ],")
        
        if obj.get('gameEffects'):
            data_lines.append(f'    gameEffects: "{escape_string(obj["gameEffects"])}",')
        
        data_lines.append("  },")
    
    data_lines.append("];")
    
    header = f"""// Auto-generated by extract_objectives.py
// DO NOT EDIT MANUALLY — regenerate from Word document

"""
    
    return header + interfaces + "\n" + "\n".join(data_lines)


def escape_string(s: str) -> str:
    """Escape a string for TypeScript string literal."""
    s = str(s)
    s = s.replace('\\', '\\\\')
    s = s.replace('"', '\\"')
    s = s.replace('\n', '\\n')
    s = s.replace('\r', '\\r')
    s = s.replace('\t', '\\t')
    return s


if __name__ == '__main__':
    docx_path = sys.argv[1] if len(sys.argv) > 1 else CORE_RULES_DOCX
    out_path = sys.argv[2] if len(sys.argv) > 2 else OUTPUT_PATH
    success = extract_objectives(docx_path, out_path)
    sys.exit(0 if success else 1)
