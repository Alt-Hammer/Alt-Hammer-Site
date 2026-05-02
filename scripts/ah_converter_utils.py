"""
ah_converter_utils.py
─────────────────────
Shared utilities for the Alt-Hammer document conversion pipeline.

This module handles the core logic that all three conversion scripts share:
  - Detecting 'Keyword' and 'Action' character styles in python-docx runs
  - Converting a paragraph (with mixed styles) into Markdown text
  - Converting Word tables into GFM markdown or HTML (for complex matrices)
  - Slugifying text for anchor links and file names
  - Writing output files safely

STYLE DETECTION
───────────────
python-docx exposes character styles on Run objects via run.style.name.
Your Word documents use two named character styles:
  - 'Keyword'  -> rendered as gold small-caps on the website
  - 'Action'   -> rendered as italic underlined on the website

Both are wrapped in <span class="keyword" data-term="slug"> on output,
with a data-type attribute to distinguish them for different tooltip styling.

TABLE HANDLING
──────────────
doc_to_markdown_sections() iterates doc.element.body directly (not
doc.paragraphs) so that <w:tbl> elements are visited in document order
alongside paragraphs. Each table is classified as:

  - MATRIX  -> hit-roll / wound-roll cross-reference tables with merged
               header cells and 10+ data columns. Rendered as a scrollable
               HTML <table class="rules-table matrix-table">.
  - SIMPLE  -> all other tables. Rendered as GFM pipe tables, which Astro
               renders natively in MDX. Cell text is sanitized for output.

A table is classified as MATRIX when it has >= MATRIX_COL_THRESHOLD columns
in its data rows (after expanding gridSpan merges). Currently 8.

IMAGE LAYOUT TABLES
───────────────────
Table 5 in the Core Rules is a 3x6 grid used purely to lay out diagram
images in a numbered 1-6 arrangement. This table has no useful text content
and is silently skipped during conversion (image embedding is out of scope).
"""

import re
import os
from pathlib import Path
from docx import Document
from docx.oxml.ns import qn

# Minimum number of data columns (after expanding gridSpan) for a table to be
# treated as a matrix and rendered as HTML rather than GFM markdown.
MATRIX_COL_THRESHOLD = 8


# ── Character sanitization ────────────────────────────────────────────────────

def sanitize_text(text: str) -> str:
    """
    Replace special Unicode characters with ASCII equivalents for MDX compatibility.
    """
    replacements = {
        '\u2013': '-',        # en dash
        '\u2014': '-',        # em dash
        '\u2018': "'",        # left single quote
        '\u2019': "'",        # right single quote
        '\u201c': '"',        # left double quote
        '\u201d': '"',        # right double quote
        '\u2026': '...',      # ellipsis
        '\u2022': '*',        # bullet
        '\u00b0': ' degrees', # degree symbol
        '\u2122': '(TM)',     # trademark
        '\u00ae': '(R)',      # registered
        '\u00a9': '(C)',      # copyright
    }

    for old, new in replacements.items():
        text = text.replace(old, new)

    return text


# ── Slug helpers ──────────────────────────────────────────────────────────────

def slugify(text: str) -> str:
    """
    Convert display text to a URL-safe slug.
    'Feel No Pain 5+' -> 'feel-no-pain-5'
    'Charge and Fight' -> 'charge-and-fight'
    """
    text = text.lower().strip()
    text = re.sub(r'[^\w\s-]', '', text)   # remove special chars except hyphens
    text = re.sub(r'[\s_]+', '-', text)     # spaces and underscores -> hyphens
    text = re.sub(r'-+', '-', text)         # collapse multiple hyphens
    return text.strip('-')


def heading_to_anchor(text: str) -> str:
    """Convert a heading to an HTML anchor id."""
    return slugify(text)


# ── Style detection ───────────────────────────────────────────────────────────

def get_run_style_name(run) -> str | None:
    """
    Return the character style name applied to a run, or None.
    Checks both run.style (explicit style) and rStyle in XML (inline style).
    """
    # Method 1: explicit style object
    if run.style and run.style.name:
        name = run.style.name
        # Exclude paragraph-level styles that bleed through
        if name not in ('Default Paragraph Style', 'Normal', 'Default'):
            return name

    # Method 2: check raw XML for rStyle element
    rpr = run._r.find(qn('w:rPr'))
    if rpr is not None:
        rstyle = rpr.find(qn('w:rStyle'))
        if rstyle is not None:
            val = rstyle.get(qn('w:val'))
            if val:
                return val

    return None


def is_keyword_run(run) -> bool:
    """Return True if this run uses the 'Keyword' character style."""
    style = get_run_style_name(run)
    if style is None:
        return False
    return style.lower() in ('keyword', 'keyword1', 'keywords', 'keywordchar', 'keyword char')


def is_action_run(run) -> bool:
    """Return True if this run uses the 'Action' character style."""
    style = get_run_style_name(run)
    if style is None:
        return False
    return style.lower() in ('action', 'action1', 'actions', 'actionchar', 'action char')


def is_weapon_run(run) -> bool:
    """Return True if this run uses the 'Weapon' character style."""
    style = get_run_style_name(run)
    if style is None:
        return False
    return style.lower() in ('weapon', 'weaponchar', 'weapon char', 'weapon ref', 'weaponref', 'weapon reference')


def is_wargear_run(run) -> bool:
    """Return True if this run uses the 'Wargear' character style."""
    style = get_run_style_name(run)
    if style is None:
        return False
    return style.lower() in ('wargear', 'wargearchar', 'wargear char', 'wargear ref', 'wargearref', 'wargear reference')


# ── Paragraph -> Markdown conversion ─────────────────────────────────────────

def runs_to_markdown(paragraph) -> str:
    """
    Convert a paragraph's runs to a Markdown string, preserving:
      - Keyword style -> <span class="keyword" data-term="slug" data-type="keyword">text</span>
      - Action style  -> <span class="keyword" data-term="slug" data-type="action">text</span>
      - Weapon style  -> <span class="weapon-ref" data-weapon="slug">text</span>
      - Wargear style -> <span class="wargear-ref" data-wargear="slug">text</span>
      - Bold          -> **text**
      - Italic        -> *text*

    Consecutive runs of the same character style are merged before emitting spans.
    This prevents Word's habit of splitting a single styled word across multiple runs
    (e.g. "M"+"ove", "Vehicles"+"s", "Anti-"+"Air") from producing broken spans.
    """
    # Step 1: classify every run
    classified = []
    for run in paragraph.runs:
        text = run.text
        if not text:
            continue
        text = sanitize_text(text)

        if is_keyword_run(run):
            classified.append(['keyword', text])
        elif is_action_run(run):
            classified.append(['action', text])
        elif is_weapon_run(run):
            classified.append(['weapon', text])
        elif is_wargear_run(run):
            classified.append(['wargear', text])
        elif run.bold and run.italic:
            classified.append(['bold-italic', text])
        elif run.bold:
            classified.append(['bold', text])
        elif run.italic:
            classified.append(['italic', text])
        else:
            classified.append(['plain', text])

    # Step 2: merge consecutive runs of the same type
    merged = []
    for item in classified:
        if merged and merged[-1][0] == item[0]:
            merged[-1][1] += item[1]
        else:
            merged.append([item[0], item[1]])

    # Step 3: emit markdown
    parts = []
    for run_type, text in merged:
        # Whitespace-only styled spans: emit as plain text
        if run_type in ('keyword', 'action', 'weapon', 'wargear') and not text.strip():
            parts.append(text)
            continue

        if run_type == 'keyword':
            slug = slugify(text.strip())
            parts.append(
                f'<span class="keyword" data-term="{slug}" data-type="keyword">{text}</span>'
            )
        elif run_type == 'action':
            slug = slugify(text.strip())
            parts.append(
                f'<span class="keyword" data-term="{slug}" data-type="action">{text}</span>'
            )
        elif run_type == 'weapon':
            slug = slugify(text.strip())
            parts.append(
                f'<span class="weapon-ref" data-weapon="{slug}">{text}</span>'
            )
        elif run_type == 'wargear':
            slug = slugify(text.strip())
            parts.append(
                f'<span class="wargear-ref" data-wargear="{slug}">{text}</span>'
            )
        elif run_type == 'bold-italic':
            parts.append(f'***{text}***')
        elif run_type == 'bold':
            parts.append(f'**{text}**')
        elif run_type == 'italic':
            parts.append(f'*{text}*')
        else:
            parts.append(text)

    return ''.join(parts)


# ── Paragraph style classification ───────────────────────────────────────────

# Maps Word heading style names to heading levels.
# Levels 1-5 are emitted as markdown (# through #####).
# Levels 6-8 are emitted as raw HTML <h6>/<h7>/<h8> tags because:
#   - Markdown only supports up to H6 (######), and H7/H8 have no syntax
#   - H6-H8 may contain tab-separated right-aligned values that need
#     <span class="heading-value"> wrapping for CSS flex layout
HEADING_STYLES = {
    'Heading 1': 1,
    'Heading 2': 2,
    'Heading 3': 3,
    'Heading 4': 4,
    'Heading 5': 5,
    'Heading 6': 6,
    'Heading 7': 7,
    'Heading 8': 8,
}

def get_heading_level(paragraph) -> int | None:
    """Return the heading level (1-7) or None if not a heading."""
    style_name = paragraph.style.name if paragraph.style else ''
    return HEADING_STYLES.get(style_name)


def is_list_paragraph(paragraph) -> bool:
    """Return True if this paragraph is a list item."""
    style_name = paragraph.style.name if paragraph.style else ''
    return 'List' in style_name or paragraph._p.find(qn('w:numPr')) is not None


def get_list_level(paragraph) -> int:
    """Return the list indentation level (0-based)."""
    numpr = paragraph._p.find(qn('w:numPr'))
    if numpr is not None:
        ilvl = numpr.find(qn('w:ilvl'))
        if ilvl is not None:
            return int(ilvl.get(qn('w:val'), 0))
    return 0


def is_numbered_list(paragraph) -> bool:
    """Return True if this is a numbered (ordered) list item."""
    style_name = paragraph.style.name if paragraph.style else ''
    return 'List Number' in style_name or 'List Paragraph' in style_name


def paragraph_to_markdown(paragraph, list_counter: dict) -> str:
    """
    Convert a single paragraph to its Markdown equivalent.

    list_counter: dict tracking current number for ordered lists per level,
                  e.g. {0: 3, 1: 1} means level-0 item 3, level-1 item 1
    """
    # Check if the entire paragraph has a Weapon or Wargear paragraph style applied.
    para_style = paragraph.style.name.lower() if paragraph.style else ''
    if para_style == 'weapon':
        raw_text = paragraph.text.strip()
        if raw_text:
            slug = slugify(raw_text)
            return f'<p><span class="weapon-ref" data-weapon="{slug}">{raw_text}</span></p>'
        return ''
    elif para_style == 'wargear':
        raw_text = paragraph.text.strip()
        if raw_text:
            slug = slugify(raw_text)
            return f'<p><span class="wargear-ref" data-wargear="{slug}">{raw_text}</span></p>'
        return ''

    text = runs_to_markdown(paragraph)

    if not text.strip():
        return ''

    heading_level = get_heading_level(paragraph)
    if heading_level is not None:
        list_counter.clear()

        if heading_level <= 5:
            # Levels 2-5: standard markdown heading syntax.
            # Tab characters are not expected here; emit as-is.
            prefix = '#' * heading_level
            return f'\n{prefix} {text}\n'

        else:
            # Levels 6-8: emit as raw HTML <h6>/<h7>/<h8>.
            # These levels may contain a tab character separating the heading
            # title from a right-aligned value (e.g. point cost, AP cost).
            # Split on the first tab and wrap the trailing part in
            # <span class="heading-value"> so the CSS flex layout can
            # push it to the right edge of the heading bar.
            tag = f'h{heading_level}'
            if '\t' in text:
                title_part, value_part = text.split('\t', 1)
                title_part = title_part.strip()
                value_part = value_part.strip()
                if value_part:
                    inner = f'{title_part}<span class="heading-value">{value_part}</span>'
                else:
                    inner = title_part
            else:
                inner = text.strip()

            # Compute an id anchor from the title text only (strip any HTML spans
            # that runs_to_markdown() may have added for keyword styling)
            import re as _re
            plain_title = _re.sub(r'<[^>]+>', '', inner.split('<span class="heading-value">')[0])
            anchor_id = slugify(plain_title)

            return f'\n<{tag} id="{anchor_id}">{inner}</{tag}>\n'

    if is_list_paragraph(paragraph):
        level = get_list_level(paragraph)
        indent = '  ' * level

        if is_numbered_list(paragraph):
            n = list_counter.get(level, 0) + 1
            list_counter[level] = n
            for k in list(list_counter.keys()):
                if k > level:
                    del list_counter[k]
            return f'{indent}{n}. {text}'
        else:
            return f'{indent}- {text}'

    list_counter.clear()
    return text


# ── Table -> Markdown / HTML conversion ──────────────────────────────────────

def _get_cell_text(cell_elem) -> str:
    """
    Extract all text from a <w:tc> XML element, joining paragraph text with
    a space so multi-paragraph cells don't lose their internal breaks entirely.
    Sanitizes special characters for MDX output.
    """
    W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    paragraphs = cell_elem.findall(f'{{{W}}}p')
    parts = []
    for p in paragraphs:
        t = ''.join(x.text or '' for x in p.iter(f'{{{W}}}t'))
        t = sanitize_text(t.strip())
        if t:
            parts.append(t)
    return ' '.join(parts)


def _get_cell_span(cell_elem) -> int:
    """Return the gridSpan value for a <w:tc> element (defaults to 1)."""
    W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    gs = cell_elem.find(f'.//{{{W}}}gridSpan')
    if gs is not None:
        val = gs.get(f'{{{W}}}val')
        if val and val.isdigit():
            return int(val)
    return 1


def _is_vmerge_continue(cell_elem) -> bool:
    """Return True if this cell is a vertical-merge continuation (no content row)."""
    W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    vm = cell_elem.find(f'.//{{{W}}}vMerge')
    if vm is None:
        return False
    val = vm.get(f'{{{W}}}val', '')
    return val != 'restart'


def _has_image(cell_elem) -> bool:
    """Return True if this cell contains an inline drawing/image."""
    ns = 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing'
    return len(cell_elem.findall(f'.//{{{ns}}}inline')) > 0


def _expand_row(row_elem) -> list:
    """
    Return an ordered list of cell text values for one <w:tr>, expanding
    gridSpan merges so each logical column has exactly one entry.
    vMerge-continue cells are represented as empty strings.
    """
    W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    cells = row_elem.findall(f'{{{W}}}tc')
    result = []
    for cell in cells:
        if _is_vmerge_continue(cell):
            result.append('')
        else:
            text = _get_cell_text(cell)
            span = _get_cell_span(cell)
            result.append(text)
            for _ in range(span - 1):
                result.append('')
    return result


def _table_is_image_layout(tbl_elem) -> bool:
    """
    Return True if this table is a pure image-layout table with no useful text.
    Heuristic: 60%+ of cells contain images or are empty.
    Used to silently skip Table 5 (the deployment-map diagram grid).
    """
    W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    rows = tbl_elem.findall(f'.//{{{W}}}tr')
    total_cells = 0
    image_or_empty_cells = 0
    for row in rows:
        cells = row.findall(f'.//{{{W}}}tc')
        for cell in cells:
            total_cells += 1
            text = _get_cell_text(cell).strip()
            if _has_image(cell) or not text:
                image_or_empty_cells += 1
    if total_cells == 0:
        return False
    return (image_or_empty_cells / total_cells) >= 0.6


def _count_data_columns(tbl_elem) -> int:
    """
    Return the maximum number of expanded columns across all data rows
    (skipping row 0, which may have spanning header cells).
    """
    W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    rows = tbl_elem.findall(f'.//{{{W}}}tr')
    max_cols = 0
    for row in rows[1:]:
        expanded = _expand_row(row)
        max_cols = max(max_cols, len(expanded))
    return max_cols


def _escape_pipe(text: str) -> str:
    """Escape pipe characters inside GFM table cells."""
    return text.replace('|', '\\|')


def table_to_gfm(tbl_elem) -> str:
    """
    Convert a simple Word table to a GFM (GitHub-Flavoured Markdown) pipe table.

    Column headers come from the first row. All subsequent rows become data rows.
    Rows that are entirely empty (vMerge continuations with no extra text) are
    skipped to keep the output clean.
    """
    W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    rows = tbl_elem.findall(f'.//{{{W}}}tr')
    if not rows:
        return ''

    grid = []
    for row in rows:
        grid.append(_expand_row(row))

    # Normalize all rows to the same column count
    max_cols = max(len(r) for r in grid) if grid else 0
    if max_cols == 0:
        return ''
    for r in grid:
        while len(r) < max_cols:
            r.append('')

    headers = [_escape_pipe(c) for c in grid[0]]
    separator = ['---'] * max_cols
    data_rows = grid[1:]

    lines = []
    lines.append('| ' + ' | '.join(headers) + ' |')
    lines.append('| ' + ' | '.join(separator) + ' |')
    for row in data_rows:
        if not any(c.strip() for c in row):
            continue
        cells = [_escape_pipe(c) for c in row]
        lines.append('| ' + ' | '.join(cells) + ' |')

    return '\n'.join(lines)


def table_to_html(tbl_elem, extra_class: str = '') -> str:
    """
    Convert a Word table to a scrollable HTML <table> block.

    Used for matrix tables (hit-roll / wound-roll cross-references) that have
    too many columns for GFM and use merged header cells.

    The first two rows are rendered as <thead> with <th> elements.
    Subsequent rows are rendered as <tbody> with <td> elements.
    gridSpan is preserved via colspan attributes.
    The table is wrapped in a <div class="table-scroll"> for horizontal scroll.
    """
    W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    rows = tbl_elem.findall(f'.//{{{W}}}tr')
    if not rows:
        return ''

    table_class = f'rules-table matrix-table {extra_class}'.strip()
    lines = [f'<div class="table-scroll">', f'<table class="{table_class}">']

    # First two rows -> thead
    lines.append('<thead>')
    for row in rows[:2]:
        cells = row.findall(f'{{{W}}}tc')
        lines.append('<tr>')
        for cell in cells:
            if _is_vmerge_continue(cell):
                continue
            text = _get_cell_text(cell)
            span = _get_cell_span(cell)
            colspan = f' colspan="{span}"' if span > 1 else ''
            safe_text = text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
            lines.append(f'  <th{colspan}>{safe_text}</th>')
        lines.append('</tr>')
    lines.append('</thead>')

    # Remaining rows -> tbody
    lines.append('<tbody>')
    for row in rows[2:]:
        cells = row.findall(f'{{{W}}}tc')
        row_cells = []
        for cell in cells:
            if _is_vmerge_continue(cell):
                continue
            text = _get_cell_text(cell)
            span = _get_cell_span(cell)
            colspan = f' colspan="{span}"' if span > 1 else ''
            safe_text = text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
            row_cells.append(f'  <td{colspan}>{safe_text}</td>')
        if row_cells:
            lines.append('<tr>')
            lines.extend(row_cells)
            lines.append('</tr>')
    lines.append('</tbody>')
    lines.append('</table>')
    lines.append('</div>')

    return '\n'.join(lines)


def table_to_markdown(tbl_elem) -> str:
    """
    Top-level dispatcher: returns the best markdown/HTML representation for a
    Word table element.

    Routing logic:
      1. Image-layout table  -> empty string (silently skipped)
      2. >= MATRIX_COL_THRESHOLD data columns -> HTML via table_to_html()
      3. Everything else     -> GFM pipe table via table_to_gfm()
    """
    if _table_is_image_layout(tbl_elem):
        return ''  # Table 5: deployment-map diagram images -- skip silently

    data_cols = _count_data_columns(tbl_elem)
    if data_cols >= MATRIX_COL_THRESHOLD:
        return table_to_html(tbl_elem)

    return table_to_gfm(tbl_elem)


# ── Document -> Markdown section conversion ───────────────────────────────────

def doc_to_markdown_sections(doc_path: str) -> list:
    """
    Parse a Word document and split it into sections by Heading 1.

    Iterates doc.element.body directly (not doc.paragraphs) so that <w:tbl>
    elements are encountered in document order alongside paragraphs. Tables
    are converted to GFM or HTML and inserted at their natural position in
    the section content.

    Returns a list of dicts:
      {
        'title': str,           # Heading 1 text
        'slug': str,            # slugified title
        'content': str,         # full Markdown content of this section
        'subsections': [str],   # list of Heading 2 titles found in section
      }
    """
    doc = Document(doc_path)

    # Build a map from XML element -> python-docx Paragraph object.
    # doc.paragraphs only covers top-level paragraphs; table-cell paragraphs
    # are NOT in this list, so para_map.get() naturally returns None for them.
    para_map = {p._element: p for p in doc.paragraphs}

    sections = []
    current_section = None
    lines = []
    list_counter = {}

    for element in doc.element.body:
        tag = element.tag.split('}')[-1]

        # ── Paragraph ────────────────────────────────────────────────────────
        if tag == 'p':
            para = para_map.get(element)
            if para is None:
                continue  # inside a table cell or other nested context

            level = get_heading_level(para)
            text = para.text.strip()

            if not text:
                if lines and lines[-1] != '':
                    lines.append('')
                continue

            if level == 1:
                # Save current section
                if current_section is not None:
                    current_section['content'] = '\n'.join(lines).strip()
                    sections.append(current_section)

                # Start new section
                current_section = {
                    'title': text,
                    'slug': slugify(text),
                    'content': '',
                    'subsections': [],
                }
                lines = []
                list_counter = {}

            else:
                if level == 2 and current_section is not None:
                    current_section['subsections'].append(text)

                md_line = paragraph_to_markdown(para, list_counter)
                if md_line:
                    is_list = is_list_paragraph(para)
                    if (lines
                        and lines[-1] != ''
                        and not lines[-1].startswith('#')
                        and not lines[-1].startswith('-')
                        and not lines[-1].startswith('  -')
                        and not is_list
                        and not any(lines[-1].startswith(f'{n}.') for n in range(1, 20))
                    ):
                        lines.append('')
                    lines.append(md_line)

        # ── Table ─────────────────────────────────────────────────────────────
        elif tag == 'tbl':
            if current_section is None:
                continue  # table before any H1 -- skip

            table_md = table_to_markdown(element)
            if table_md:
                if lines and lines[-1] != '':
                    lines.append('')
                lines.append(table_md)
                lines.append('')

    # Save last section
    if current_section is not None:
        current_section['content'] = '\n'.join(lines).strip()
        sections.append(current_section)

    return sections


# ── File output helpers ───────────────────────────────────────────────────────

def ensure_dir(path: str):
    """Create directory if it doesn't exist."""
    Path(path).mkdir(parents=True, exist_ok=True)


def write_mdx_file(output_path: str, frontmatter: dict, content: str):
    """
    Write a .mdx file with YAML frontmatter and Markdown content.

    frontmatter: dict of key/value pairs written as YAML
    content: Markdown body text
    """
    ensure_dir(os.path.dirname(output_path))

    yaml_lines = ['---']
    for key, value in frontmatter.items():
        if isinstance(value, str):
            escaped = value.replace('"', '\\"')
            yaml_lines.append(f'{key}: "{escaped}"')
        elif isinstance(value, list):
            yaml_lines.append(f'{key}:')
            for item in value:
                yaml_lines.append(f'  - "{item}"')
        else:
            yaml_lines.append(f'{key}: {value}')
    yaml_lines.append('---')
    yaml_lines.append('')

    full_content = '\n'.join(yaml_lines) + '\n' + content + '\n'

    with open(output_path, 'w', encoding='utf-8') as f:
        f.write(full_content)

    print(f'  ✓  Written: {output_path}')


def write_json_file(output_path: str, data):
    """Write a JSON file with pretty formatting."""
    import json
    ensure_dir(os.path.dirname(output_path))
    with open(output_path, 'w', encoding='utf-8') as f:
        json.dump(data, f, indent=2, ensure_ascii=False)
    print(f'  ✓  Written: {output_path}')
